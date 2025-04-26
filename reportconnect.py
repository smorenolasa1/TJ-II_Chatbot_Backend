from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from datetime import datetime
import os, json, textwrap
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Pt, Inches
from dotenv import load_dotenv
import google.generativeai as genai
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

# Configuración inicial
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
MODEL_NAME = "models/gemini-1.5-pro"

app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

CONTEXT_DIR = "context"
STATIC_DIR = "static"
os.makedirs(CONTEXT_DIR, exist_ok=True)
os.makedirs(STATIC_DIR, exist_ok=True)
def is_markdown_table(lines):
    return (
        len(lines) >= 2
        and "|" in lines[0]
        and "|" in lines[1]
        and set(lines[1].strip()) <= set("|- ")
    )
@app.get("/generate_report")
def generate_report():
    files = {
        "SimilPatternTool": "similpattern_history.json",
        "ShotLlama2": "shotllama2_history.json",
        "CsvUpdate": "csvupdate_history.json"
    }

    sections = []
    for section, file in files.items():
        path = os.path.join(CONTEXT_DIR, file)
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                entries = json.load(f)
                if entries:
                    summary = f"### {section}\n"
                    for entry in entries:
                        summary += f"**Query:** {entry.get('question', '')}\n"
                        if "pattern_summary" in entry:
                            summary += f"**Pattern Summary:**\n{entry['pattern_summary']}\n"
                        if "response" in entry:
                            summary += f"**Results:**\n{entry['response']}\n"
                        if "plot_path" in entry:
                            summary += f"**Plot:** {entry['plot_path']}\n"
                        summary += "\n"
                    sections.append(summary)

    if not sections:
        return JSONResponse({"message": "No context to generate report."}, status_code=200)

    report_text = "# Plasma Fusion Tools Report\n\n" + "\n".join(sections)

    prompt = f"""
You are a scientific assistant generating a clean, structured report based on plasma fusion tools analysis.

Generate clear section headers (no #), concise bullet points if useful, and format numerical tables when needed.

DO NOT write meta-instructions like "Here's the report".

Don´t write the similarity pattern info like confidence shot and interval, since it is already being displayed in a graph.
Do write some small analysis of the results.
Structure it like this when possible:

Tool Name
Query: ...
Pattern Summary:
[Table]
Results: ...
Plot: [keep the relative plot path]

Here is the raw input:
{report_text}
"""

    model = genai.GenerativeModel(MODEL_NAME)
    response = model.generate_content(prompt)
    cleaned = response.text.strip()

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    pdf_filename = f"report_{timestamp}.pdf"
    docx_filename = f"report_{timestamp}.docx"
    pdf_path = os.path.join(STATIC_DIR, pdf_filename)
    docx_path = os.path.join(STATIC_DIR, docx_filename)

    # ➤ DOCX
    doc = Document()
    doc.add_heading("Fusion Data Analysis Report", 0)

    for block in cleaned.split("\n\n"):
        lines = block.strip().split("\n")
        if not lines:
            continue

        if lines[0].lower().startswith("plot:"):
            plot_path = lines[0].split(":", 1)[1].strip()
            full_path = os.path.join(plot_path)
            if os.path.exists(full_path):
                doc.add_picture(full_path, width=Inches(5.5))
        elif is_markdown_table(lines):
            # Parse markdown table to 2D list
            table_data = [
                [cell.strip() for cell in row.strip("|").split("|")]
                for row in lines
                if not re.match(r"^\s*\|?\s*-+\s*\|", row)  # Skip separator line
            ]
            table = doc.add_table(rows=0, cols=len(table_data[0]))
            table.style = 'Table Grid'

            for row_data in table_data:
                row = table.add_row().cells
                for idx, cell in enumerate(row_data):
                    row[idx].text = cell
        else:
            for line in lines:
                if not line.lower().startswith("query:"):
                    doc.add_paragraph(line.strip(), style='Normal')

    doc.save(docx_path)

    # ➤ PDF
    pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    story = []

    for block in cleaned.split("\n\n"):
        lines = block.strip().split("\n")
        if not lines:
            continue

        if lines[0].lower().startswith("plot:"):
            plot_path = lines[0].split(":", 1)[1].strip()
            full_path = os.path.join(plot_path)
            if os.path.exists(full_path):
                story.append(Spacer(1, 12))
                img = Image(full_path, width=500, height=200)
                story.append(img)
                story.append(Spacer(1, 24))
        elif all(line.strip().startswith("|") and line.strip().endswith("|") for line in lines):
            # Table detected
            data = [
                line.strip("|").split("|")
                for line in lines
                if not all(cell.strip("- ") == "" for cell in line.strip("|").split("|"))  # descarta fila de --- separadora
            ]
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#4B72B0")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
            ]))
            story.append(table)
            story.append(Spacer(1, 16))
        else:
            for line in lines:
                if not line.lower().startswith("query:"):
                    story.append(Paragraph(line.strip(), styles["Normal"]))
            story.append(Spacer(1, 12))

    pdf_doc.build(story)

    print("✅ Returning report URLs:")
    print(f"PDF: http://localhost:5005/static/{pdf_filename}")
    print(f"Word: http://localhost:5005/static/{docx_filename}")

    return {
        "pdf_url": f"http://localhost:5005/static/{pdf_filename}",
        "word_url": f"http://localhost:5005/static/{docx_filename}"
    }

@app.post("/reset_context")
def reset_context():
    deleted = []
    for file in ["similpattern_history.json", "shotllama2_history.json", "csvupdate_history.json"]:
        path = os.path.join(CONTEXT_DIR, file)
        if os.path.exists(path):
            os.remove(path)
            deleted.append(file)
    return {"message": f"Deleted: {deleted}"}

app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")